import os
import yaml
import json
import uuid
import shutil
import re
import time
import asyncio
import logging
from pathlib import Path
from google import genai
from google.genai import types
from pypdf import PdfReader
import pandas as pd
import docx
from dotenv import load_dotenv
import warnings

# --- KONFIGURATION FÖR TURBO MODE ---
INITIAL_CONCURRENCY = 5
MAX_CONCURRENCY = 50
MIN_CONCURRENCY = 1
GAS_STEP = 1.0
BRAKE_FACTOR = 0.5
COOLDOWN_TIME = 10
WASHER_CHUNK_SIZE = 25000

# --- LOGGNING ---
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    datefmt='%H:%M:%S'
)
log = logging.getLogger("Pipeline")

# --- CLEANER PROMPT ---
CLEANER_PROMPT_TEMPLATE = """
Du är en "data-tvätt"-specialist. Ditt mål är att extrahera den rena brödtexten.

INSTRUKTIONER:
1. Identifiera semantisk brödtext (regler, instruktioner, tabellinnehåll).
2. RADERA ALLT metadata-brus:
    * Sidhuvuden/Sidfötter (t.ex. "Adda Inköpscentral", "Sida 1 av 5")
    * Sidnummer
    * Dokument-IDn som upprepas på varje sida
3. Behåll strukturen (Rubriker, listor) intakt.
4. Svara ENDAST med den rena texten. Inget prat.

--- TEXT ATT TVÄTTA ---
{raw_text_chunk}
"""

# --- 1. SETUP ---

CURRENT_DIR = Path(__file__).parent.resolve()
PROJECT_ROOT = CURRENT_DIR.parent.parent 

def load_yaml(path_relative_to_root):
    full_path = PROJECT_ROOT / path_relative_to_root
    if not full_path.exists():
        raise FileNotFoundError(f"CRITICAL: Saknas: {full_path}")
    with open(full_path, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)

def load_text(path_relative_to_root):
    full_path = PROJECT_ROOT / path_relative_to_root
    if not full_path.exists():
        raise FileNotFoundError(f"CRITICAL: Saknas: {full_path}")
    with open(full_path, "r", encoding="utf-8") as f:
        return f.read()

# INITIERING
try:
    log.info(f"🔧 Initierar Pipeline v6.5 (Full Format Support) från: {PROJECT_ROOT}")
    
    env_path = PROJECT_ROOT / "ai-services" / ".env"
    if env_path.exists():
        load_dotenv(env_path)
    
    CONFIG_PATH = "ai-services/data_pipeline/config/pipeline_config.yaml"
    PIPELINE_CONFIG = load_yaml(CONFIG_PATH)
    
    ctx_path = PIPELINE_CONFIG['paths'].get('dir_master_context')
    if not ctx_path:
        ctx_path = "ai-services/data_pipeline/config/master_context_protocol.md"
    
    CONTEXT_PROTOCOL = load_text(ctx_path)

except Exception as e:
    log.error(f"Initieringsfel: {e}")
    exit(1)

api_key = os.environ.get("GOOGLE_API_KEY")
if not api_key:
    log.error("GOOGLE_API_KEY saknas.")
    exit(1)

client = genai.Client(api_key=api_key)

# --- 2. ADAPTIVE THROTTLER ---

class AdaptiveThrottler:
    def __init__(self, start, min_limit, max_limit):
        self.limit = float(start)
        self.min_limit = min_limit
        self.max_limit = max_limit
        self.current_running = 0
        self.lock = asyncio.Lock()
        self.last_brake = 0

    async def wait_for_slot(self):
        while True:
            async with self.lock:
                if self.current_running < int(self.limit):
                    self.current_running += 1
                    return
            await asyncio.sleep(0.1)

    def release(self):
        self.current_running -= 1

    async def gas(self):
        async with self.lock:
            if self.limit < self.max_limit:
                self.limit += GAS_STEP

    async def brake(self):
        async with self.lock:
            now = time.time()
            if now - self.last_brake < 2.0: return 
            self.last_brake = now
            old_limit = self.limit
            self.limit = max(self.min_limit, self.limit * BRAKE_FACTOR)
            log.warning(f"🛑 429 DETECTED! Bromsar: {int(old_limit)} -> {int(self.limit)}. Paus {COOLDOWN_TIME}s.")
        await asyncio.sleep(COOLDOWN_TIME)

# --- 3. I/O & TEXT EXTRACTION ---

def extract_text_from_pdf(pdf_path):
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            c = page.extract_text()
            if c: text += c + "\n"
        return text
    except Exception as e:
        log.error(f"PDF Read Error {pdf_path.name}: {e}")
        return None

def extract_text_from_excel(file_path):
    try:
        warnings.filterwarnings('ignore', category=UserWarning, module='openpyxl')
        dfs = pd.read_excel(file_path, sheet_name=None)
        text = ""
        for sheet, df in dfs.items():
            text += f"\n--- SHEET: {sheet} ---\n"
            text += df.head(100).to_markdown(index=False)
            text += "\n"
        return text
    except Exception as e:
        log.error(f"Excel Read Error {file_path.name}: {e}")
        return None

def extract_text_from_csv(file_path):
    """NYTT: Hanterar CSV-filer via Pandas."""
    try:
        # Försök gissa separator (vanligtvis , eller ;)
        try:
            df = pd.read_csv(file_path, sep=None, engine='python')
        except:
            df = pd.read_csv(file_path) # Fallback till default
            
        text = f"\n--- CSV DATA: {file_path.name} ---\n"
        text += df.head(100).to_markdown(index=False)
        return text
    except Exception as e:
        log.error(f"CSV Read Error {file_path.name}: {e}")
        return None

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                full_text.append(" | ".join(row_data))
            full_text.append("")
        return "\n".join(full_text)
    except Exception as e:
        log.error(f"DOCX Read Error {file_path.name}: {e}")
        return None

def extract_plain_text(file_path):
    """NYTT: Hanterar .txt och .md filer."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        log.error(f"Text Read Error {file_path.name}: {e}")
        return None

# --- 4. FORMATTERING & TAGGAR ---

def sanitize_tag_for_obsidian(tag):
    tag = str(tag).lower().strip()
    tag = re.sub(r'[\s-]', '_', tag)
    tag = re.sub(r'[^\w_]', '', tag)
    if not tag: return None
    if tag.isdigit(): tag = f"n{tag}"
    return tag

def generate_tags_from_filename(filename):
    clean_name = Path(filename).stem
    clean_name = re.sub(r'[-_]', ' ', clean_name)
    words = set(word.lower() for word in clean_name.split() if len(word) > 2)
    return list(words)

def save_full_document_reference(text, original_filename, zone):
    block_id = str(uuid.uuid4())
    auto_tags = generate_tags_from_filename(original_filename)
    auto_tags.append("full_document")
    
    clean_tags = []
    for t in auto_tags:
        safe_tag = sanitize_tag_for_obsidian(t)
        if safe_tag: clean_tags.append(safe_tag)
    
    yaml_header = "---\n"
    yaml_header += f"uuid: \"{block_id}\"\n"
    yaml_header += f"doc_type: \"smart_block\"\n"
    yaml_header += f"source_file: \"{original_filename}\"\n"
    yaml_header += f"authority_level: \"{zone}\"\n"
    yaml_header += f"block_type: \"FULL_DOCUMENT\"\n"
    yaml_header += f"process_step: [\"general\"]\n"
    yaml_header += f"tags: {json.dumps(clean_tags)}\n"
    yaml_header += "---\n\n"
    
    full_content = yaml_header + text
    safe_name = Path(original_filename).stem[:30].replace(" ", "_")
    filename = f"general_FULL_DOCUMENT_{zone.upper()}_{safe_name}_{block_id[:4]}.md"
    
    output_path = PROJECT_ROOT / PIPELINE_CONFIG['paths']['dir_converted']
    output_path.mkdir(parents=True, exist_ok=True)
    
    with open(output_path / filename, "w", encoding="utf-8") as f:
        f.write(full_content)

def save_smart_block(block_data, original_filename, zone):
    block_id = str(uuid.uuid4())
    meta = block_data.get("metadata", {})
    
    content_raw = block_data.get("content_markdown")
    if not content_raw or not isinstance(content_raw, str) or not content_raw.strip():
        return 
    
    block_type = meta.get('block_type', 'DEFINITION')
    steps = meta.get('process_step', ['general'])
    
    raw_tags = meta.get('tags', [])
    clean_tags = []
    for t in raw_tags:
        safe_tag = sanitize_tag_for_obsidian(t)
        if safe_tag: clean_tags.append(safe_tag)

    yaml_header = "---\n"
    yaml_header += f"uuid: \"{block_id}\"\n"
    yaml_header += f"doc_type: \"smart_block\"\n"
    yaml_header += f"source_file: \"{original_filename}\"\n"
    yaml_header += f"authority_level: \"{zone}\"\n"
    yaml_header += f"block_type: \"{block_type}\"\n"
    yaml_header += f"process_step: {json.dumps(steps)}\n"
    yaml_header += f"tags: {json.dumps(clean_tags)}\n"
    
    if "graph_relations" in meta:
        yaml_header += f"graph_relations: {json.dumps(meta['graph_relations'])}\n"
    if "constraints" in meta:
        yaml_header += f"constraints: {json.dumps(meta['constraints'])}\n"
    yaml_header += "---\n\n"
    
    full_content = yaml_header + content_raw
    safe_step = steps[0] if steps else "general"
    # Extrahera bara siffran från step (t.ex. "step_1_intake" -> "1_intake")
    step_short = safe_step.replace("step_", "") if safe_step.startswith("step_") else safe_step
    filename = f"{step_short}_{block_type}_{zone.upper()}_{block_id[:8]}.md"
    
    output_path = PROJECT_ROOT / PIPELINE_CONFIG['paths']['dir_converted']
    output_path.mkdir(parents=True, exist_ok=True)
    
    with open(output_path / filename, "w", encoding="utf-8") as f:
        f.write(full_content)

# --- 5. ASYNKRON AI-LOGIK ---

def clean_json_response(text):
    cleaned = re.sub(r'^```json\s*', '', text, flags=re.MULTILINE)
    cleaned = re.sub(r'^```\s*', '', cleaned, flags=re.MULTILINE)
    return cleaned.strip()

async def wash_text_async(raw_text, filename, throttler):
    chunks = []
    for i in range(0, len(raw_text), WASHER_CHUNK_SIZE):
        chunks.append(raw_text[i : i + WASHER_CHUNK_SIZE])
    
    cleaned_chunks = []
    model_name = PIPELINE_CONFIG['system']['models']['model_fast'] 
    
    for i, chunk in enumerate(chunks):
        prompt = CLEANER_PROMPT_TEMPLATE.format(raw_text_chunk=chunk)
        
        while True:
            await throttler.wait_for_slot()
            try:
                response = await client.aio.models.generate_content(
                    model=model_name,
                    contents=prompt
                )
                await throttler.gas()
                throttler.release()
                
                cleaned_text = response.text.strip() if response.text else ""
                cleaned_chunks.append(cleaned_text)
                break 

            except Exception as e:
                err_str = str(e)
                if "429" in err_str or "ResourceExhausted" in err_str:
                    throttler.release()
                    await throttler.brake()
                    continue
                else:
                    throttler.release()
                    log.error(f"Washer Error {filename} chunk {i}: {e}")
                    cleaned_chunks.append(chunk) 
                    break
    
    return "\n\n".join(cleaned_chunks)

async def analyze_document_async(text, zone, filename, throttler):
    model_name = PIPELINE_CONFIG['system']['models']['model_pro']
    
    prompt = f"""
{CONTEXT_PROTOCOL}

--- ANALYSINSTRUKTION ---
FILNAMN: {filename}
ZON: {zone}

DU SKA: Returnera en strikt JSON-lista med smarta block.
VIKTIGT: Om texten innehåller mätbara regler (volym, tid, pengar, regioner) MÅSTE du fylla i 'constraints'-listan i metadata enligt protokollet (Sektion 8).

TEXT INPUT:
{text[:60000]}
"""
    
    while True:
        await throttler.wait_for_slot()
        try:
            response = await client.aio.models.generate_content(
                model=model_name,
                contents=prompt,
                config=types.GenerateContentConfig(response_mime_type="application/json")
            )
            await throttler.gas()
            throttler.release()
            
            cleaned = clean_json_response(response.text)
            data = json.loads(cleaned)
            if isinstance(data, dict):
                if "blocks" in data: return data["blocks"]
                return [data]
            return data

        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "ResourceExhausted" in err_str:
                throttler.release()
                await throttler.brake()
                continue
            else:
                throttler.release()
                log.error(f"Analysis Error {filename}: {e}")
                return []

# --- 6. WORKER ---

async def process_file_concurrently(file_path, zone, throttler):
    # Dynamisk textläsning baserat på filändelse
    text = None
    ext = file_path.suffix.lower()
    
    if ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif ext == '.xlsx':
        text = extract_text_from_excel(file_path)
    elif ext == '.docx':
        text = extract_text_from_docx(file_path)
    elif ext == '.csv':
        text = extract_text_from_csv(file_path)
    elif ext in ['.txt', '.md']:
        text = extract_plain_text(file_path)
    
    if not text or len(text.strip()) == 0:
        return

    # Tvätt
    cleaned_text = await wash_text_async(text, file_path.name, throttler)
    
    # Spara Helhet
    save_full_document_reference(cleaned_text, file_path.name, zone)
    
    # Analys
    log.info(f"⏳ Analys köas: {file_path.name} (Active: {throttler.current_running}/{int(throttler.limit)})")
    
    blocks = await analyze_document_async(cleaned_text, zone, file_path.name, throttler)
    
    if blocks:
        for block in blocks:
            save_smart_block(block, file_path.name, zone)
        
        archive_rel = PIPELINE_CONFIG['paths'].get('dir_archive')
        if archive_rel:
            archive_path = PROJECT_ROOT / archive_rel
            archive_path.mkdir(parents=True, exist_ok=True)
            shutil.move(str(file_path), str(archive_path / file_path.name))
        
        log.info(f"✅ Klar: {file_path.name} -> {len(blocks)} block.")
    else:
        log.warning(f"⚠️  Inga block för: {file_path.name}")

# --- 7. MAIN ---

async def main():
    log.info("--- ADDA INGEST PIPELINE v6.5 (FULL FORMAT SUPPORT) ---")
    
    throttler = AdaptiveThrottler(INITIAL_CONCURRENCY, MIN_CONCURRENCY, MAX_CONCURRENCY)
    tasks = []
    
    # Hämta tillåtna extensions från config
    allowed_exts = PIPELINE_CONFIG.get('processing', {}).get('allowed_extensions', ['.pdf'])
    allowed_exts = [e.lower() for e in allowed_exts]
    log.info(f"Tillåtna format: {allowed_exts}")

    for zone, dir_key in [("PRIMARY", 'dir_primary_raw'), ("SECONDARY", 'dir_secondary_raw')]:
        path = PROJECT_ROOT / PIPELINE_CONFIG['paths'][dir_key]
        if path.exists():
            all_files = []
            for ext in allowed_exts:
                all_files.extend(path.glob(f"*{ext}"))
            
            for f in all_files:
                tasks.append(process_file_concurrently(f, zone, throttler))
            
    if not tasks:
        log.warning("Inga filer att bearbeta.")
        return
        
    log.info(f"🚀 Startar {len(tasks)} jobb...")
    start_time = time.time()
    await asyncio.gather(*tasks)
    log.info(f"✨ Klara på {time.time() - start_time:.2f} sekunder.")

if __name__ == "__main__":
    asyncio.run(main())